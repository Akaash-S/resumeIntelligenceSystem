import io
import hashlib
from loguru import logger
import PyPDF2
import pdfplumber
import docx
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

def get_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()

def parse_pdf_pypdf2(file_bytes: bytes) -> str:
    logger.bind(stage="PARSE").debug("Attempting PDF extraction with PyPDF2...")
    pdf_file = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def parse_pdf_pdfplumber(file_bytes: bytes) -> str:
    logger.bind(stage="PARSE").debug("Attempting PDF extraction with pdfplumber...")
    pdf_file = io.BytesIO(file_bytes)
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def parse_pdf_ocr(file_bytes: bytes) -> str:
    logger.bind(stage="PARSE").info("Starting PDF OCR fallback via PyMuPDF and Tesseract...")
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            logger.bind(stage="PARSE").debug(f"OCRing page {page_num + 1}/{len(doc)}...")
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        logger.bind(stage="PARSE").error(f"OCR failed or Tesseract is not installed properly: {e}")
        return ""

def parse_pdf(file_bytes: bytes) -> str:
    # 1. Try PyPDF2
    text = ""
    try:
        text = parse_pdf_pypdf2(file_bytes)
    except Exception as e:
        logger.bind(stage="PARSE").warning(f"PyPDF2 failed: {e}")

    # 2. Try pdfplumber if text is empty/too short
    if len(text.strip()) < 100:
        logger.bind(stage="PARSE").info("Text from PyPDF2 is too short, falling back to pdfplumber...")
        try:
            plumber_text = parse_pdf_pdfplumber(file_bytes)
            if len(plumber_text.strip()) > len(text.strip()):
                text = plumber_text
        except Exception as e:
            logger.bind(stage="PARSE").warning(f"pdfplumber failed: {e}")

    # 3. Try OCR if still empty/too short
    if len(text.strip()) < 100:
        logger.bind(stage="PARSE").info("Text from parsing is too short (< 100 chars), triggering OCR pipeline...")
        text = parse_pdf_ocr(file_bytes)

    return text

def parse_docx(file_bytes: bytes) -> str:
    logger.bind(stage="PARSE").debug("Attempting DOCX extraction...")
    doc_file = io.BytesIO(file_bytes)
    doc = docx.Document(doc_file)
    text = []
    for para in doc.paragraphs:
        if para.text:
            text.append(para.text)
            
    # Also parse tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text.append(" | ".join(row_text))
                
    return "\n".join(text).strip()

def parse_txt(file_bytes: bytes) -> str:
    logger.bind(stage="PARSE").debug("Attempting plain text extraction...")
    decodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
    for encoding in decodings:
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    # If all fail, decode with ignore error replacement
    return file_bytes.decode("utf-8", errors="ignore").strip()

def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.split(".")[-1].lower()
    logger.bind(stage="PARSE").info(f"Extracting text from {filename} ({ext})...")
    
    if ext == "pdf":
        text = parse_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        text = parse_docx(file_bytes)
    elif ext in ["txt", "md"]:
        text = parse_txt(file_bytes)
    else:
        logger.bind(stage="PARSE").error(f"Unsupported file format: {ext}")
        raise ValueError(f"Unsupported file format: {ext}")
        
    if not text.strip():
        logger.bind(stage="PARSE").warning(f"Extracted text for {filename} is empty.")
        raise ValueError("Extracted text is empty or unparseable.")
        
    logger.bind(stage="PARSE").info(f"Successfully extracted {len(text)} characters from {filename}")
    return text
